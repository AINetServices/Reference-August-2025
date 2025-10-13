"""
Custom tools for resume processing and text extraction.
"""

import os
import io
from typing import Any, Dict, List, Type
from langchain.tools import BaseTool
from langchain_groq import ChatGroq
import PyPDF2
import docx
from pydantic import BaseModel, Field


# Define input schemas for each tool
class ResumeParserInput(BaseModel):
    file_path: str = Field(description="Path to the resume file to parse")


class ReferenceExtractorInput(BaseModel):
    resume_text: str = Field(description="The resume text to extract references from")
    role: str = Field(default="", description="Target role for reference extraction")


class ResumeAnalyzerInput(BaseModel):
    resume_text: str = Field(description="The resume text to analyze")
    target_role: str = Field(default="", description="Target role for analysis")


# Helper functions
def _extract_pdf_text(file_content: io.BytesIO) -> str:
    """Extract text from PDF file content (BytesIO)"""
    try:
        file_content.seek(0)  # Reset to beginning
        pdf_reader = PyPDF2.PdfReader(file_content)
        text = ""
        
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return text.strip()
    except Exception as e:
        return f"Error reading PDF: {str(e)}"


def _extract_docx_text(file_content: io.BytesIO) -> str:
    """Extract text from DOCX file content (BytesIO)"""
    try:
        file_content.seek(0)  # Reset to beginning
        doc = docx.Document(file_content)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"


def _parse_llm_response(response: str) -> Dict[str, Any]:
    """Parse LLM response into structured format"""
    try:
        import json
        
        # Find JSON in the response
        start_idx = response.find('{')
        end_idx = response.rfind('}') + 1
        
        if start_idx >= 0 and end_idx > start_idx:
            json_str = response[start_idx:end_idx]
            return json.loads(json_str)
        else:
            return {
                "references": [],
                "total_found": 0,
                "extraction_notes": "Failed to parse LLM response",
                "raw_response": response
            }
    except json.JSONDecodeError:
        return {
            "references": [],
            "total_found": 0,
            "extraction_notes": "Invalid JSON in LLM response",
            "raw_response": response
        }


def _parse_analysis_response(response: str) -> Dict[str, Any]:
    """Parse comprehensive analysis response"""
    try:
        import json
        
        start_idx = response.find('{')
        end_idx = response.rfind('}') + 1
        
        if start_idx >= 0 and end_idx > start_idx:
            json_str = response[start_idx:end_idx]
            return json.loads(json_str)
        else:
            return {
                "error": "Failed to parse analysis response",
                "raw_response": response
            }
    except json.JSONDecodeError:
        return {
            "error": "Invalid JSON in analysis response",
            "raw_response": response
        }


# Create simple function-based tools
def create_resume_parser_tool(llm: ChatGroq):
    """Create a resume parser tool"""
    
    def parse_resume(file_path: str) -> str:
        """Parse a resume file and return extracted text"""
        try:
            with open(file_path, 'rb') as file:
                filename = file_path.lower()
                if filename.endswith('.pdf'):
                    return _extract_pdf_text(file)
                elif filename.endswith(('.docx', '.doc')):
                    return _extract_docx_text(file)
                else:
                    return "Unsupported file format"
        except Exception as e:
            return f"Error parsing file: {str(e)}"
    
    return parse_resume


def create_reference_extractor_tool(llm: ChatGroq):
    """Create a reference extractor tool"""
    
    def extract_references(resume_text: str, role: str = "") -> Dict[str, Any]:
        """Extract references from resume text"""
        try:
            prompt = f"""
            Analyze this resume text and extract professional reference information:

            Resume Text:
            {resume_text}

            Target Role: {role}

            Extract and identify references and return as JSON.
            """
            
            response = llm.invoke(prompt)
            return _parse_llm_response(response.content)
        except Exception as e:
            return {"error": f"Failed to extract references: {str(e)}"}
    
    return extract_references


def create_resume_analyzer_tool(llm: ChatGroq):
    """Create a resume analyzer tool"""
    
    def analyze_resume(resume_text: str, target_role: str = "") -> Dict[str, Any]:
        """Analyze resume content and extract comprehensive insights"""
        try:
            analysis_prompt = f"""
            Perform a comprehensive analysis of this resume for the target role: {target_role}

            Resume Text:
            {resume_text}

            Analyze and return as structured JSON.
            """
            
            response = llm.invoke(analysis_prompt)
            return _parse_analysis_response(response.content)
        except Exception as e:
            return {"error": f"Failed to analyze resume: {str(e)}"}
    
    return analyze_resume


# Simple wrapper classes
class ResumeParserTool:
    def __init__(self, llm: ChatGroq):
        self.llm = llm
        self._parse_func = create_resume_parser_tool(llm)
    
    def _run(self, file_path: str) -> str:
        return self._parse_func(file_path)
    
    def parse_file_content(self, file_content: io.BytesIO, filename: str) -> str:
        """
        Parse resume content from file bytes - this is what the workflow expects
        """
        try:
            # Since your current parser expects a file path, we need to adapt it
            # Save the bytes to a temporary file and use your existing parser
            import tempfile
            import os
            
            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp_file:
                # Write the bytes content to the temp file
                file_content.seek(0)  # Reset to beginning
                tmp_file.write(file_content.read())
                tmp_file_path = tmp_file.name
            
            # Use your existing parser
            result = self._parse_func(tmp_file_path)
            
            # Clean up the temp file
            os.unlink(tmp_file_path)
            
            return result
            
        except Exception as e:
            return f"Error parsing file content: {str(e)}"
    
    def parse_resume(self, file_content: io.BytesIO, filename: str) -> str:
        """Alternative method name that calls parse_file_content"""
        return self.parse_file_content(file_content, filename)


class ReferenceExtractorTool:
    def __init__(self, llm: ChatGroq):
        self.llm = llm
        self._extract_func = create_reference_extractor_tool(llm)
    
    def _run(self, resume_text: str, role: str = "") -> Dict[str, Any]:
        return self._extract_func(resume_text, role)


class ResumeAnalyzerTool:
    def __init__(self, llm: ChatGroq):
        self.llm = llm
        self._analyze_func = create_resume_analyzer_tool(llm)
    
    def _run(self, resume_text: str, target_role: str = "") -> Dict[str, Any]:
        return self._analyze_func(resume_text, target_role)